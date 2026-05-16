from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx

doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_paragraph(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table_row(table, col1, col2, col3, col4, bold=False):
    row = table.add_row()
    cells = row.cells
    for i, text in enumerate([col1, col2, col3, col4]):
        cells[i].text = text
        for para in cells[i].paragraphs:
            for run in para.runs:
                run.bold = bold
                run.font.size = Pt(11)
                run.font.name = 'Times New Roman'


# ==================== 3.2 SYSTEM ARCHITECTURE ====================
add_heading(doc, "3.2 System Architecture", level=1)

add_paragraph(doc,
    "The architecture of PaddyScan was planned keeping simplicity and scalability in mind. "
    "Since this is a mobile-first application dealing with image processing and machine learning, "
    "we needed a structure that cleanly separates the user interface from the business logic and "
    "the actual data handling. After going through a few options, we settled on a three-layered "
    "architecture that works well with Flutter's ecosystem and also gives us room to expand later "
    "if needed."
)

add_paragraph(doc,
    "The three main layers are the Presentation Layer, the Application Layer, and the Data and "
    "Intelligence Layer. Each layer has its own responsibilities and they communicate through "
    "well-defined boundaries. Below is a breakdown of each layer and how they contribute to the "
    "overall system."
)

# 3.2.1 Presentation Layer
add_heading(doc, "3.2.1 Presentation Layer", level=2)

add_paragraph(doc,
    "This is essentially everything the user sees and interacts with. The Presentation Layer in "
    "PaddyScan is built entirely using Flutter widgets and screens. We have six main screens in "
    "the app: a splash screen that shows on startup, the home page where users can pick or capture "
    "an image, a results page that shows the disease analysis output, a history page that lists "
    "previous scans, a settings page where users can configure things like server IP and language, "
    "and a full-screen image viewer."
)

add_paragraph(doc,
    "For state management, we chose the BLoC (Business Logic Component) pattern using the "
    "flutter_bloc package. The main BLoC here is the HomeBloc, which handles everything from "
    "image picking to triggering the analysis and receiving the results. The reason we picked BLoC "
    "over something simpler like setState is that it keeps the UI completely separate from the "
    "logic, which made debugging a lot easier during development."
)

add_paragraph(doc,
    "Navigation is handled using GoRouter, which allows us to define routes declaratively and also "
    "supports deep linking. The bottom navigation shell (MainShell) wraps the three main sections: "
    "Scan, History, and Settings, so users can switch between them easily without losing their state."
)

# 3.2.2 Application Layer
add_heading(doc, "3.2.2 Application Layer", level=2)

add_paragraph(doc,
    "The Application Layer sits between the UI and the data. It contains the business logic that "
    "decides what should happen when a user performs an action. In PaddyScan, this layer includes "
    "the use cases and the BLoC event handlers."
)

add_paragraph(doc,
    "When a user uploads an image, the PickImageUseCase is responsible for picking the image, "
    "compressing it (we resize to a max of 1024px to speed up uploads), and converting it to "
    "base64 format so it can be sent to the Flask backend. The UploadImageUseCase then takes "
    "that data and passes it to the appropriate API endpoint depending on the selected analysis "
    "mode — classify, detect, or diagnose."
)

add_paragraph(doc,
    "This layer also handles error scenarios. We built a custom exception hierarchy so that errors "
    "like network timeouts, server errors, or image picking cancellations each get their own "
    "meaningful message rather than a generic crash. This made testing and user feedback much "
    "more manageable."
)

add_paragraph(doc,
    "Dependency injection is set up using the GetIt service locator, which means each component "
    "gets what it needs passed in rather than creating dependencies internally. This keeps things "
    "testable and modular."
)

# 3.2.3 Data and Intelligence Layer
add_heading(doc, "3.2.3 Data and Intelligence Layer", level=2)

add_paragraph(doc,
    "This is where the actual processing happens. The Data and Intelligence Layer is split into "
    "two parts: the local data handling on the Flutter side, and the backend intelligence that "
    "runs on a Python Flask server."
)

add_paragraph(doc,
    "On the Flutter side, data persistence is managed through SharedPreferences and platform-aware "
    "file storage. The HistoryService saves each scan result locally, including the image and the "
    "prediction details. On mobile, images are saved as actual files in the app's cache directory. "
    "On web, they are stored as base64 strings. The app keeps a maximum of 50 history entries and "
    "automatically removes older ones."
)

add_paragraph(doc,
    "The intelligence part lives in the Flask backend. It uses two machine learning models: a "
    "TensorFlow classification model (.h5 file) that identifies the disease type and confidence "
    "score, and a YOLOv8 PyTorch detection model (.pt file) that finds the exact locations of "
    "infected regions in the image. Both models are loaded once when the Flask server starts up "
    "to avoid repeated loading delays."
)

add_paragraph(doc,
    "The backend also does image annotation — it draws bounding boxes over affected areas and "
    "calculates what percentage of the rice plant is diseased. This processed image is saved and "
    "a URL is returned to the Flutter app so it can display the annotated version to the user."
)

add_paragraph(doc,
    "Additionally, the app supports AI-powered text elaboration through an Ollama/Deepseek model "
    "integration and also has the Google Gemini API configured for generating disease explanations. "
    "Offline, it falls back to a hardcoded disease information database stored in disease_info.dart."
)

# 3.2.4 Data Flow and Communication
add_heading(doc, "3.2.4 Data Flow and Communication", level=2)

add_paragraph(doc,
    "Understanding how data moves through the system helped us catch a few design issues early "
    "on. The flow starts the moment a user picks an image from their camera or gallery."
)

add_paragraph(doc,
    "Step 1 — Image Capture: The user taps either the camera or gallery button on the home page. "
    "The HomeBloc receives a PickImageEvent, calls the image_picker package, and once the image "
    "is selected, it is compressed asynchronously on a separate thread. The resulting bytes are "
    "base64-encoded and stored in an ImageData object."
)

add_paragraph(doc,
    "Step 2 — Mode Selection: The user picks one of the three analysis modes. This is just a UI "
    "state change that updates the selectedMode in HomeState."
)

add_paragraph(doc,
    "Step 3 — Upload and Analysis: When the user taps Analyze, the HomeBloc fires an "
    "AnalyzeImageEvent. The ApiService reads the current server IP from SharedPreferences "
    "(so if the user changed it in settings, it picks it up automatically), then sends a POST "
    "request using Dio to the correct Flask endpoint — /api/classify, /api/detect, or "
    "/api/diagnose."
)

add_paragraph(doc,
    "Step 4 — Response Handling: The Flask server processes the image, runs the appropriate ML "
    "models, and returns a JSON response with the disease label, confidence, bounding box "
    "coordinates (if detection mode), affected area percentage, and a URL to the annotated image. "
    "The Flutter app parses this into a PredictionResult model."
)

add_paragraph(doc,
    "Step 5 — Result Display and History Save: The HomeBloc emits a success state and the UI "
    "navigates to the ResultPage. Simultaneously, the result is saved to local history via "
    "HistoryService so the user can review it later."
)

add_paragraph(doc,
    "All communication between Flutter and Flask uses HTTP with JSON payloads. There is no "
    "WebSocket or real-time streaming for the image analysis — it is a simple request-response "
    "cycle with a 45-second timeout configured in Dio."
)

# ==================== 3.3 USE CASE ANALYSIS ====================
add_heading(doc, "3.3 Use Case Analysis", level=1)

add_paragraph(doc,
    "This section describes the main use cases of PaddyScan from the end user's perspective. "
    "Each use case outlines what triggers it, who is involved, what steps happen, and what the "
    "expected outcome is. The goal here is to make sure the system actually does what a farmer "
    "or agricultural worker would expect it to do in real situations."
)

# Helper to add use case table
def add_use_case(doc, uc_id, title, actors, preconditions, main_flow, alt_flow, postconditions):
    add_heading(doc, title, level=2)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.5)

    def add_row(label, content):
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = content
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
        row.cells[0].paragraphs[0].runs[0].bold = True

    add_row("Use Case ID", uc_id)
    add_row("Use Case Name", title.split(": ", 1)[-1] if ": " in title else title)
    add_row("Actor(s)", actors)
    add_row("Preconditions", preconditions)
    add_row("Main Flow", main_flow)
    add_row("Alternative Flow", alt_flow)
    add_row("Postconditions", postconditions)

    doc.add_paragraph()

# UC 3.3.1
add_heading(doc, "3.3.1 Use Case: Navigate to Different Pages", level=2)
add_paragraph(doc,
    "Use Case ID: UC-01"
)
add_paragraph(doc,
    "Actor(s): Application User (farmer or agricultural worker)"
)
add_paragraph(doc,
    "Preconditions: The PaddyScan application is installed and successfully launched on the user's device."
)
add_paragraph(doc,
    "Description: This use case covers how a user moves between the three main sections of the app — "
    "the Scan page, the History page, and the Settings page. Navigation is done through a bottom "
    "navigation bar that is always visible once the splash screen finishes."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The user opens the app and the splash screen plays for a couple of seconds.",
    "2. The app lands on the Home (Scan) page by default.",
    "3. The user taps the History icon at the bottom to go to the scan history.",
    "4. The user taps Settings to configure server IP, theme, or language.",
    "5. The user taps Scan to return to the main scanning page.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
add_paragraph(doc,
    "- If the app is launched for the very first time, shared preferences are empty and the "
    "app uses default settings (English language, dark theme, default server IP).",
    indent=True
)
add_paragraph(doc, "Postconditions: The selected page is shown and the user's current state is preserved during navigation.")

doc.add_paragraph()

# UC 3.3.2
add_heading(doc, "3.3.2 Use Case: Upload an Image for Disease Detection", level=2)
add_paragraph(doc, "Use Case ID: UC-02")
add_paragraph(doc, "Actor(s): Application User")
add_paragraph(doc,
    "Preconditions: The user is on the Home/Scan page and the Flask backend server is running and reachable."
)
add_paragraph(doc,
    "Description: The user needs to provide an image of a rice plant so the app can analyze it. "
    "They can either capture a new photo using the device camera or choose an existing image from "
    "their gallery. This is the first step in the disease detection process."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The user taps the camera button to take a live photo, or the gallery button to select an existing one.",
    "2. The device asks for camera or storage permission if not already granted.",
    "3. The user captures or selects the image.",
    "4. The app compresses the image (resizes to max 1024px) and shows a preview on the home page.",
    "5. The user selects an analysis mode (Classify, Detect, or Diagnose).",
    "6. The user taps the Analyze button to send the image to the server.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
alts = [
    "- If the user denies camera/gallery permission, an error message is shown and the image is not picked.",
    "- If the user cancels the image picker, the app stays on the home page with no image loaded.",
    "- If the image is not a valid format (not JPEG/PNG), a validation error is shown.",
]
for a in alts:
    add_paragraph(doc, a, indent=True)

add_paragraph(doc, "Postconditions: A valid image is loaded and displayed on the home page, ready for analysis.")

doc.add_paragraph()

# UC 3.3.3
add_heading(doc, "3.3.3 Use Case: Analyze Uploaded Image for Rice Disease", level=2)
add_paragraph(doc, "Use Case ID: UC-03")
add_paragraph(doc, "Actor(s): Application User, Flask Backend Server, ML Models")
add_paragraph(doc,
    "Preconditions: An image has already been selected or captured. The server is connected (green indicator on home page)."
)
add_paragraph(doc,
    "Description: Once the image is uploaded, the core feature of the app kicks in. The image "
    "is sent to the Flask server which runs it through the trained machine learning models. "
    "Depending on the selected mode, the system either classifies the disease, detects its location, "
    "or does a full diagnosis combining both."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The HomeBloc receives the AnalyzeImageEvent after the user taps Analyze.",
    "2. The image (in base64 format) is sent via a POST request to the appropriate Flask endpoint.",
    "3. The Flask server receives the image, decodes it, and passes it to the ML model(s).",
    "4. For Classify mode: the TensorFlow model outputs a disease label and confidence percentage.",
    "5. For Detect mode: the YOLOv8 model finds affected regions, draws bounding boxes, and computes affected area percentage.",
    "6. For Diagnose mode: both models run and their outputs are combined into a full report.",
    "7. The backend returns a JSON response with the results and a URL to the annotated image.",
    "8. The Flutter app parses the response and navigates to the Results page.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
alts = [
    "- If the server is unreachable or times out (45s), a NetworkException is raised and the user sees an error message.",
    "- If the server returns an error status, a ServerException is raised with the HTTP status code.",
    "- If the JSON response cannot be parsed, a DataFormatException is raised.",
]
for a in alts:
    add_paragraph(doc, a, indent=True)

add_paragraph(doc, "Postconditions: The Results page displays the disease name, confidence score, annotated image, affected percentage, and an AI-generated explanation.")

doc.add_paragraph()

# UC 3.3.4
add_heading(doc, "3.3.4 Use Case: Delete Uploaded Image", level=2)
add_paragraph(doc, "Use Case ID: UC-04")
add_paragraph(doc, "Actor(s): Application User")
add_paragraph(doc, "Preconditions: The user is on the History page and there is at least one saved scan in the list.")
add_paragraph(doc,
    "Description: Users may want to clean up their scan history to free storage or remove "
    "irrelevant scans. This use case allows them to delete individual records or wipe the entire history."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The user navigates to the History page from the bottom navigation bar.",
    "2. The list of previous scans is loaded from local SharedPreferences storage.",
    "3. The user long-presses or swipes a scan entry to reveal the delete option.",
    "4. The user confirms the deletion.",
    "5. The HistoryService removes the entry from SharedPreferences and also deletes the locally saved image file.",
    "6. The list refreshes and the deleted entry no longer appears.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
alts = [
    "- The user can tap 'Clear All' to delete all history entries at once. A confirmation dialog appears before this action is executed.",
    "- If the history is already empty, the page shows an empty state message.",
]
for a in alts:
    add_paragraph(doc, a, indent=True)

add_paragraph(doc, "Postconditions: The selected scan record(s) are permanently deleted from the device. The history list reflects the updated state.")

doc.add_paragraph()

# UC 3.3.5
add_heading(doc, "3.3.5 Use Case: Access Cloud Service Links for Image Upload", level=2)
add_paragraph(doc, "Use Case ID: UC-05")
add_paragraph(doc, "Actor(s): Application User")
add_paragraph(doc, "Preconditions: The app is running and the user is connected to the internet or local network.")
add_paragraph(doc,
    "Description: In some cases, especially on web or when the device cannot directly connect to the "
    "local Flask server, the user may need to access cloud-hosted service links for uploading images. "
    "This use case covers how users configure or access the server endpoint and any cloud service URLs "
    "available through the Settings page."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The user navigates to the Settings page.",
    "2. The user finds the 'Server Configuration' section.",
    "3. The current server IP address and port are shown (default: 192.168.1.101:5000).",
    "4. The user enters the cloud service URL or updated IP address in the text field.",
    "5. The new address is saved to SharedPreferences.",
    "6. From this point on, all API calls use the updated server address.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
alts = [
    "- If the entered URL is unreachable, the health check on the home page will show a red/disconnected status.",
    "- The user can reset to the default IP if they make a mistake.",
]
for a in alts:
    add_paragraph(doc, a, indent=True)

add_paragraph(doc, "Postconditions: The new server/cloud link is saved and will be used for all subsequent image upload requests.")

doc.add_paragraph()

# UC 3.3.6
add_heading(doc, "3.3.6 Use Case: View Team and Department Information", level=2)
add_paragraph(doc, "Use Case ID: UC-06")
add_paragraph(doc, "Actor(s): Application User")
add_paragraph(doc, "Preconditions: The app is installed and running.")
add_paragraph(doc,
    "Description: Users may want to know who built the app or what institution or department is "
    "behind it. This use case provides access to the About or Team information section within the app."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The user navigates to the Settings page.",
    "2. The user scrolls to the 'About' or 'Team' section.",
    "3. The team names, department, and institution details are displayed.",
    "4. Any relevant contact or project information is shown.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
add_paragraph(doc, "- No alternative flows. This is a read-only, static information screen.", indent=True)
add_paragraph(doc, "Postconditions: The user has viewed the team and department information. No system state changes.")

doc.add_paragraph()

# UC 3.3.7
add_heading(doc, "3.3.7 Use Case: Access Precautionary Measures and Sources", level=2)
add_paragraph(doc, "Use Case ID: UC-07")
add_paragraph(doc, "Actor(s): Application User")
add_paragraph(doc, "Preconditions: A disease has been detected and the Results page is currently displayed, or the user is viewing a past scan from history.")
add_paragraph(doc,
    "Description: After a disease is identified, the app does not just show a label — it also "
    "provides the user with actionable information. This includes precautionary measures, treatment "
    "suggestions, and links to reliable agricultural sources. This feature is especially useful for "
    "farmers who may not have immediate access to an agronomist."
)
add_paragraph(doc, "Main Flow:")
steps = [
    "1. The Results page shows the detected disease name and confidence.",
    "2. Below the result summary, an AI-generated explanation is displayed using a typing effect.",
    "3. The explanation includes what the disease is, how it spreads, and how to manage it.",
    "4. If the AI model is unavailable (no network), the app falls back to the local disease_info.dart database.",
    "5. The information is rendered as formatted Markdown text using the flutter_markdown package.",
    "6. Source references or links to further reading are shown at the bottom of the explanation.",
]
for s in steps:
    add_paragraph(doc, s, indent=True)

add_paragraph(doc, "Alternative Flow:")
alts = [
    "- If the AI elaboration API call fails, the app silently falls back to the offline disease info without showing an error.",
    "- If the detected disease has no offline entry in disease_info.dart, a generic message is shown instead.",
]
for a in alts:
    add_paragraph(doc, a, indent=True)

add_paragraph(doc, "Postconditions: The user has read the precautionary measures and relevant sources for the detected disease. No changes are made to app state.")

# Save
output_path = r"e:\Flutter\Projects\paddyscan\PaddyScan_Architecture_UseCases.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
